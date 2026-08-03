from __future__ import annotations

import asyncio
import json
import os
import subprocess
import sys
from contextlib import nullcontext
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock

import pytest
from fastapi import HTTPException
from pydantic import ValidationError
from starlette.datastructures import UploadFile

from src import config_loader, lightrag_service, model_profiles
from src.api import server
from src.doc_processor.parsers.base_parser import Document
from src.exceptions import ManifestCorruptedError
from src.lightrag_service import LightRAGService
from src.runtime_lock import RuntimeLock


def _service(tmp_path: Path, workspace: str = "kb") -> LightRAGService:
    return LightRAGService(
        {
            "paths": {
                "data_dir": str(tmp_path / "data"),
                "lightrag_dir": str(tmp_path / "lightrag"),
            },
            "lightrag": {"workspace": "default"},
            "siliconflow": {},
        },
        workspace=workspace,
    )


def test_manifest_restores_valid_backup_and_refuses_double_corruption(tmp_path):
    service = _service(tmp_path)
    service._save_manifest({"documents": {"doc_old": {"doc_name": "old.txt"}}})
    service._save_manifest({"documents": {"doc_new": {"doc_name": "new.txt"}}})

    service.manifest_path.write_text("{broken", encoding="utf-8")
    restored = service._load_manifest()

    assert set(restored["documents"]) == {"doc_old"}
    assert json.loads(service.manifest_path.read_text(encoding="utf-8"))["documents"]

    service.manifest_path.write_text("{broken", encoding="utf-8")
    service.manifest_backup_path.write_text("{also-broken", encoding="utf-8")
    with pytest.raises(ManifestCorruptedError, match="backup"):
        service._load_manifest()


def test_failed_reindex_keeps_active_version_and_marks_attempt_stale(
    tmp_path,
    monkeypatch,
):
    service = _service(tmp_path)
    doc = Document(
        doc_id="source",
        file_name="example.txt",
        file_path=str(tmp_path / "example.txt"),
        file_type="txt",
        raw_text="new content",
        metadata={"lightrag_doc_id": "doc_123"},
    )
    service._save_manifest(
        {
            "documents": {
                "doc_123": {
                    "doc_id": "doc_123",
                    "doc_name": "example.txt",
                    "file_type": "txt",
                    "file_path": doc.file_path,
                    "indexed": True,
                    "status": "processed",
                    "active_index_status": "processed",
                    "active_index_doc_id": "doc_123",
                    "content_sha256": "old",
                }
            }
        }
    )
    fake_rag = SimpleNamespace(addon_params={})
    service._rag = fake_rag
    monkeypatch.setattr(service, "assert_embedding_compatible", lambda: None)
    monkeypatch.setattr(service, "cleanup_interrupted_index_docs", AsyncMock(return_value=[]))
    monkeypatch.setattr(
        service,
        "_temporary_index_llm_and_kg_filter",
        lambda *_args, **_kwargs: nullcontext(),
    )
    monkeypatch.setattr(
        service,
        "_insert_document_text",
        AsyncMock(side_effect=RuntimeError("provider failed")),
    )
    monkeypatch.setattr(service, "discard_lightrag_doc", AsyncMock())

    class Collector:
        on_update = None

        def scope(self):
            return nullcontext()

        def to_stages(self):
            return {}

    monkeypatch.setattr(lightrag_service, "install_stage_timing", lambda _rag: Collector())

    with pytest.raises(RuntimeError, match="provider failed"):
        asyncio.run(service.index_document(doc))

    item = service._load_manifest()["documents"]["doc_123"]
    assert item["indexed"] is True
    assert item["status"] == "processed"
    assert item["active_index_doc_id"] == "doc_123"
    assert item["index_stale"] is True
    assert item["last_index_attempt_status"] == "failed"
    assert item["last_index_attempt"]["index_doc_id"].startswith("doc_123-v")


def test_shadow_commit_rolls_back_workspace_manifest_and_embedding_meta(
    tmp_path,
    monkeypatch,
):
    active = _service(tmp_path, "kb")
    active.workspace_dir.mkdir(parents=True)
    (active.workspace_dir / "state.txt").write_text("old", encoding="utf-8")
    active._save_manifest({"documents": {"old": {"doc_name": "old.txt"}}})
    active.record_embedding_signature(
        {"base_url": "https://old.test/v1", "model": "old", "embed_dim": 3},
        overwrite=True,
    )

    shadow_config = {
        "paths": {
            "data_dir": str(tmp_path / "data"),
            "lightrag_dir": str(tmp_path / "shadow" / "lightrag"),
            "lightrag_manifest_override": str(tmp_path / "shadow" / "manifest.json"),
            "lightrag_embedding_meta_dir": str(tmp_path / "shadow" / "embedding"),
        },
        "lightrag": {"workspace": "default"},
        "siliconflow": {},
    }
    shadow = LightRAGService(shadow_config, workspace="kb")
    shadow.workspace_dir.mkdir(parents=True)
    (shadow.workspace_dir / "state.txt").write_text("new", encoding="utf-8")
    shadow._save_manifest({"documents": {"new": {"doc_name": "new.txt"}}})
    shadow.record_embedding_signature(
        {"base_url": "https://new.test/v1", "model": "new", "embed_dim": 4},
        overwrite=True,
    )

    monkeypatch.setattr(server, "get_lightrag_service", lambda _workspace: active)
    monkeypatch.setattr(server, "reset_lightrag_service_async", AsyncMock())
    real_replace = os.replace

    def fail_on_shadow_embedding(src, dst):
        if Path(src).resolve() == shadow.embedding_meta_path.resolve():
            raise OSError("simulated metadata swap failure")
        return real_replace(src, dst)

    monkeypatch.setattr(server.os, "replace", fail_on_shadow_embedding)

    with pytest.raises(OSError, match="metadata swap"):
        asyncio.run(
            server._commit_shadow_rebuild(
                {"workspace": "kb", "task_id": "task"},
                shadow,
            )
        )

    assert (active.workspace_dir / "state.txt").read_text(encoding="utf-8") == "old"
    assert set(active._load_manifest()["documents"]) == {"old"}
    meta = json.loads(active.embedding_meta_path.read_text(encoding="utf-8"))
    assert meta["model"] == "old"


def test_model_mutations_are_blocked_while_index_task_is_active(monkeypatch):
    monkeypatch.setattr(
        server,
        "_index_tasks",
        {"task": {"task_id": "task", "status": "running"}},
    )
    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            server.api_upsert_model_profile(
                server.ModelProfileRequest(
                    name="provider",
                    api_base="https://api.siliconflow.cn/v1",
                )
            )
        )
    assert exc.value.status_code == 409
    assert exc.value.detail["code"] == "MODEL_CONFIG_BUSY"

    with pytest.raises(HTTPException) as config_exc:
        asyncio.run(server.update_model_config(server.ModelConfig()))
    assert config_exc.value.status_code == 409


def test_default_profile_cannot_be_deleted_and_bound_profile_falls_back(tmp_path):
    config = {
        "paths": {"data_dir": str(tmp_path / "data")},
        "siliconflow": {},
    }
    with pytest.raises(ValueError, match="不能删除"):
        model_profiles.delete_profile(model_profiles.DEFAULT_PROFILE_ID, config)

    created = model_profiles.upsert_profile(
        {
            "id": "custom",
            "name": "Custom",
            "api_base": "https://api.siliconflow.cn/v1",
        },
        config,
    )
    assert created["id"] == "custom"
    model_profiles.save_bindings(
        {"chat": {"profile_id": "custom", "model": "custom-model"}},
        config,
    )

    result = model_profiles.delete_profile("custom", config)
    bindings = model_profiles.get_bindings(config)
    defaults = model_profiles._default_store(config)["bindings"]

    assert result["binding_fallbacks"]["chat"] == defaults["chat"]
    assert bindings["chat"] == defaults["chat"]


def test_runtime_lock_rejects_second_process(tmp_path):
    lock_path = tmp_path / "runtime.lock"
    with RuntimeLock(lock_path):
        script = (
            "from src.runtime_lock import RuntimeLock\n"
            "from src.exceptions import RuntimeLockError\n"
            f"p = r'''{lock_path}'''\n"
            "try:\n"
            "    RuntimeLock(p).acquire()\n"
            "except RuntimeLockError:\n"
            "    raise SystemExit(23)\n"
            "raise SystemExit(0)\n"
        )
        result = subprocess.run(
            [sys.executable, "-c", script],
            cwd=Path(__file__).resolve().parents[1],
            check=False,
            capture_output=True,
            text=True,
        )
    assert result.returncode == 23


def test_retrieval_failure_is_not_reported_as_empty_context(monkeypatch):
    class BrokenService:
        async def preview_context(self, *_args, **_kwargs):
            raise RuntimeError("vector backend unavailable")

    monkeypatch.setattr(server, "_ensure_workspace_available", lambda _workspace: None)
    monkeypatch.setattr(server, "_ensure_embedding_compatible", lambda _workspace: None)
    monkeypatch.setattr(server, "get_lightrag_service", lambda _workspace: BrokenService())

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            server.recall_test(
                server.RecallTestRequest(workspace="kb", query="question")
            )
        )
    assert exc.value.status_code == 502
    assert exc.value.detail["code"] == "RAG_RETRIEVAL_FAILED"


def test_session_turn_lock_serializes_and_releases_references():
    order: list[str] = []

    async def worker(name: str, delay: float):
        async with server._session_turn("abc123abc123"):
            order.append(f"{name}:start")
            await asyncio.sleep(delay)
            order.append(f"{name}:end")

    async def run():
        await asyncio.gather(worker("a", 0.02), worker("b", 0))

    asyncio.run(run())

    assert order == ["a:start", "a:end", "b:start", "b:end"]
    assert "abc123abc123" not in server._session_turn_locks
    assert "abc123abc123" not in server._session_turn_lock_refs


def test_explicit_missing_session_is_not_recreated(monkeypatch):
    monkeypatch.setattr(server, "_ensure_workspace_available", lambda _workspace: None)
    monkeypatch.setattr(server, "_ensure_embedding_compatible", lambda _workspace: None)
    monkeypatch.setattr(server, "get_lightrag_service", lambda _workspace: object())
    monkeypatch.setattr(server, "_load_session", lambda _session_id: None)

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            server._chat_send_impl(
                server.ChatSendRequest(
                    session_id="abc123abc123",
                    workspace="kb",
                    message="hello",
                )
            )
        )
    assert exc.value.status_code == 404


def test_stream_retrieval_failure_returns_structured_error_and_releases_lock(
    monkeypatch,
):
    session_id = "abc123abc124"
    failure = HTTPException(
        status_code=502,
        detail={
            "code": "RAG_RETRIEVAL_FAILED",
            "detail": "vector backend unavailable",
        },
    )
    monkeypatch.setattr(
        server,
        "_chat_send_stream_impl",
        AsyncMock(side_effect=failure),
    )

    async def run():
        response = await server.chat_send_stream(
            server.ChatSendRequest(
                session_id=session_id,
                workspace="kb",
                message="question",
            )
        )
        chunks = []
        async for chunk in response.body_iterator:
            chunks.append(chunk.decode() if isinstance(chunk, bytes) else chunk)
        return "".join(chunks)

    body = asyncio.run(run())

    assert "event: error" in body
    payload = json.loads(body.split("data: ", 1)[1].strip())
    assert payload["code"] == "RAG_RETRIEVAL_FAILED"
    assert payload["detail"] == "vector backend unavailable"
    assert payload["request_id"]
    assert session_id not in server._session_turn_locks
    assert session_id not in server._session_turn_lock_refs


def test_upload_limit_stops_stream_and_removes_partial_file(tmp_path):
    upload = UploadFile(filename="large.txt", file=BytesIO(b"x" * 11))
    destination = tmp_path / "upload.tmp"

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            server._stream_upload_to_file(
                upload,
                destination,
                max_bytes=10,
            )
        )
    assert exc.value.status_code == 413


def test_batch_and_workspace_validation_reject_bad_inputs():
    with pytest.raises(ValidationError):
        server.BatchIndexRequest(workspace="kb", doc_names=[])
    with pytest.raises(ValidationError):
        server.WorkspaceCreateRequest(workspace="../escape")


def test_cli_preflight_rejects_recursive_duplicate_basenames(tmp_path):
    from src.app.cli import _preflight_documents

    (tmp_path / "a").mkdir()
    (tmp_path / "b").mkdir()
    (tmp_path / "a" / "same.txt").write_text("one", encoding="utf-8")
    (tmp_path / "b" / "same.txt").write_text("two", encoding="utf-8")

    with pytest.raises(ValueError, match="Duplicate document basenames"):
        _preflight_documents(tmp_path)


def test_cli_rebuild_preflight_rejects_unmanaged_source_changes(tmp_path, monkeypatch):
    from src.app import cli

    data_dir = tmp_path / "data"
    source_dir = tmp_path / "source"
    upload_dir = data_dir / "uploads" / "kb"
    source_dir.mkdir()
    upload_dir.mkdir(parents=True)
    (source_dir / "same.txt").write_text("new", encoding="utf-8")
    (upload_dir / "same.txt").write_text("old", encoding="utf-8")
    service = LightRAGService(
        {
            "paths": {
                "data_dir": str(data_dir),
                "lightrag_dir": str(data_dir / "lightrag"),
            },
            "lightrag": {"workspace": "default"},
            "siliconflow": {},
        },
        workspace="kb",
    )
    service._save_manifest(
        {"documents": {"doc": {"doc_name": "same.txt", "indexed": True}}}
    )
    monkeypatch.setattr(
        cli,
        "get_config",
        lambda: {
            "paths": {
                "data_dir": str(data_dir),
                "lightrag_dir": str(data_dir / "lightrag"),
            },
            "lightrag": {"workspace": "default"},
            "siliconflow": {},
        },
    )

    with pytest.raises(ValueError, match="ingest changes first"):
        cli._preflight_rebuild_sources(source_dir, "kb")


def test_effective_config_write_path_uses_override_or_local(tmp_path, monkeypatch):
    custom = tmp_path / "custom.yaml"
    monkeypatch.setenv("LIGHTGRAPHRAG_CONFIG_PATH", str(custom))
    assert config_loader.get_effective_write_config_path() == custom

    monkeypatch.delenv("LIGHTGRAPHRAG_CONFIG_PATH")
    local = tmp_path / "config" / "local.yaml"
    monkeypatch.setenv("LIGHTGRAPHRAG_CONFIG_LOCAL_PATH", str(local))
    assert config_loader.get_effective_write_config_path() == local


def test_supported_document_parsers_and_loader_deduplication(tmp_path):
    import docx
    import fitz

    from src.doc_processor.loader import DocumentLoader

    markdown_path = tmp_path / "guide.md"
    markdown_path.write_text("# Guide\n\nBody with `code`.", encoding="utf-8")

    html_path = tmp_path / "page.html"
    html_path.write_text(
        "<html><head><title>Page</title><meta name='author' content='Tester'>"
        "<script>ignore()</script></head><body><header>nav</header><p>HTML body</p></body></html>",
        encoding="utf-8",
    )

    docx_path = tmp_path / "manual.docx"
    word = docx.Document()
    word.core_properties.title = "Manual"
    word.add_paragraph("DOCX body")
    table = word.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "left"
    table.cell(0, 1).text = "right"
    word.save(docx_path)

    pdf_path = tmp_path / "paper.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "PDF body")
    pdf.set_metadata({"title": "Paper", "author": "Tester"})
    pdf.save(pdf_path)
    pdf.close()

    duplicate_path = tmp_path / "duplicate.md"
    duplicate_path.write_bytes(markdown_path.read_bytes())

    loader = DocumentLoader()
    parsed = {
        path.suffix: loader.load_document(path)
        for path in (markdown_path, html_path, docx_path, pdf_path)
    }

    assert parsed[".md"].metadata["title"] == "Guide"
    assert "HTML body" in parsed[".html"].raw_text
    assert "ignore" not in parsed[".html"].raw_text
    assert "left | right" in parsed[".docx"].raw_text
    assert "PDF body" in parsed[".pdf"].raw_text
    assert len(loader.load_all(tmp_path)) == 4
