"""DOCX document parser using python-docx."""

from __future__ import annotations

from pathlib import Path

import docx
from loguru import logger

from src.doc_processor.parsers.base_parser import BaseParser, Document


class DocxParser(BaseParser):
    """Parser for .docx files using python-docx."""

    SUPPORTED_EXTENSIONS = [".docx"]

    def parse(self, file_path: Path) -> Document:
        """Parse a .docx file into a Document object.

        Args:
            file_path: Path to the .docx file.

        Returns:
            Document with extracted text and metadata.
        """
        try:
            doc = docx.Document(str(file_path))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            raw_text = "\n\n".join(paragraphs)

            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        tables_text.append(row_text)
            if tables_text:
                raw_text += "\n\n--- Tables ---\n" + "\n".join(tables_text)

            # Metadata from core properties
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
            }

            doc_id = f"docx_{file_path.parent.name}_{file_path.stem}"
            logger.info(f"Parsed DOCX: {file_path.name}, {len(paragraphs)} paragraphs, {len(doc.tables)} tables")
            return Document(
                doc_id=doc_id,
                file_name=file_path.name,
                file_path=str(file_path),
                file_type="docx",
                raw_text=raw_text,
                metadata=metadata,
            )
        except Exception as e:
            logger.error(f"Failed to parse DOCX {file_path}: {e}")
            raise
